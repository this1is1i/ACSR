"""将论文初稿.md转换为格式化的.docx文件（不需要pandoc）"""
import re
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in kwargs.items():
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), val.get('val', 'single'))
        element.set(qn('w:sz'), val.get('sz', '4'))
        element.set(qn('w:color'), val.get('color', '000000'))
        tcBorders.append(element)
    tcPr.append(tcBorders)

def add_table_borders(table):
    """为表格所有单元格添加边框"""
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={'val': 'single', 'sz': '4'},
                bottom={'val': 'single', 'sz': '4'},
                left={'val': 'single', 'sz': '4'},
                right={'val': 'single', 'sz': '4'})

def parse_markdown(filepath):
    """解析Markdown文件，返回结构化内容列表"""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    elements = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks (```...```)
        if line.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith('```'):
                code_lines.append(lines[i].rstrip())
                i += 1
            elements.append(('code', '\n'.join(code_lines)))
            i += 1
            continue

        # Math blocks ($$...$$)
        if line.startswith('$$'):
            math_lines = []
            i += 1
            math_content = line[2:].strip()
            if math_content:
                math_lines.append(math_content)
            while i < len(lines) and not lines[i].startswith('$$'):
                math_lines.append(lines[i].rstrip())
                i += 1
            # Skip closing $$ but keep its content
            closing = lines[i] if i < len(lines) else '$$'
            if closing.startswith('$$') and len(closing) > 2:
                extra = closing[2:].strip()
                if extra:
                    math_lines.append(extra)
            i += 1
            elements.append(('math', '\n'.join(math_lines)))
            continue

        # Headings
        if line.startswith('# ') or line.startswith('## ') or line.startswith('### ') or line.startswith('#### '):
            level = len(re.match(r'^#+', line).group())
            title = line[level:].strip()
            elements.append(('heading', title, level))
            i += 1
            continue

        # Horizontal rule
        if line.strip() == '---':
            elements.append(('hr', None))
            i += 1
            continue

        # Table detection (line with | ... |)
        if line.strip().startswith('|') and line.strip().endswith('|'):
            table_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].rstrip())
                i += 1
            # Skip separator line (|---|---|)
            cleaned = [l for l in table_lines if not re.match(r'^\|[\s\-:|]+\|$', l.strip())]
            if len(cleaned) >= 1:
                rows = []
                for tl in cleaned:
                    cells = [c.strip() for c in tl.strip().split('|')[1:-1]]
                    rows.append(cells)
                elements.append(('table', rows))
            continue

        # Empty line
        if not line.strip():
            i += 1
            continue

        # Regular paragraph (accumulate until empty line)
        para_lines = [line.strip()]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].strip().startswith('#') and not lines[i].strip().startswith('```') and not lines[i].strip().startswith('$$') and not lines[i].strip().startswith('|') and lines[i].strip() != '---':
            para_lines.append(lines[i].strip())
            i += 1
        elements.append(('para', ' '.join(para_lines)))

    return elements

def apply_inline_formatting(run, text):
    """处理行内格式：**粗体**、*斜体*、`代码`、[链接](url)"""
    # Handle bold and italic
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            sub = run.add_run()
            sub.text = part[2:-2]
            sub.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            sub = run.add_run()
            sub.text = part[1:-1]
            sub.italic = True
        else:
            # Clean up HTML tags and links
            cleaned = re.sub(r'<[^>]+>', '', part)
            cleaned = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', cleaned)
            cleaned = cleaned.replace('`', '')
            run.add_run().text = cleaned

def create_docx(elements, output_path):
    """从解析后的元素创建Word文档"""
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(12)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置标题样式
    for level, (name, size) in {1: ('Heading 1', 18), 2: ('Heading 2', 15), 3: ('Heading 3', 13)}.items():
        try:
            h_style = doc.styles[name]
            h_font = h_style.font
            h_font.name = '黑体'
            h_font.size = Pt(size)
            h_font.bold = True
            h_font.color.rgb = RGBColor(0, 0, 0)
            h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            h_style.paragraph_format.space_before = Pt(12)
            h_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass

    # 页码
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    for elem in elements:
        if elem[0] == 'heading':
            _, title, level = elem
            level = min(level, 3)  # Max heading level 3
            p = doc.add_heading(title, level=level)
            # Ensure Chinese font for headings
            for run in p.runs:
                run.font.name = '黑体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
                run.font.color.rgb = RGBColor(0, 0, 0)

        elif elem[0] == 'para':
            text = elem[1]
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Pt(24)  # 两个字符缩进
            p.paragraph_format.line_spacing = 1.5
            run = p.add_run()
            # Bold prefix handling (如 **关键词：**)
            bold_match = re.match(r'^(\*\*(.+?)\*\*)', text)
            if bold_match:
                bold_run = p.add_run()
                bold_run.text = bold_match.group(2)
                bold_run.bold = True
                remaining = text[len(bold_match.group(1)):]
                if remaining:
                    apply_inline_formatting(p, remaining)
            else:
                apply_inline_formatting(p, text)

        elif elem[0] == 'code':
            code_text = elem[1]
            # Code as formatted paragraph
            for line in code_text.split('\n'):
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.line_spacing = 1.0
                run = p.add_run(line)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)

        elif elem[0] == 'math':
            math_text = elem[1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(math_text)
            run.italic = True
            run.font.size = Pt(11)

        elif elem[0] == 'table':
            rows = elem[1]
            if not rows:
                continue
            num_cols = max(len(r) for r in rows)
            # Normalize rows
            for r in rows:
                while len(r) < num_cols:
                    r.append('')
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            add_table_borders(table)
            for ri, row_data in enumerate(rows):
                for ci, cell_text in enumerate(row_data):
                    cell = table.rows[ri].cells[ci]
                    cell.text = ''
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.font.size = Pt(10)
                    if ri == 0:
                        run.bold = True
            # Add spacing after table
            doc.add_paragraph()

        elif elem[0] == 'hr':
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run('—' * 40)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(180, 180, 180)

    doc.save(output_path)
    print(f'DOCX saved to: {output_path}')

if __name__ == '__main__':
    import sys
    input_file = 'docs/论文初稿.md'
    output_file = 'docs/论文初稿.docx'

    print(f'Parsing {input_file}...')
    elements = parse_markdown(input_file)
    print(f'Found {len(elements)} elements')

    print(f'Creating {output_file}...')
    create_docx(elements, output_file)
